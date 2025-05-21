import io
import logging
from PyPDF2 import PdfReader
import docx # python-docx

# Configure basic logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_text_from_document(file_path: str, file_content: bytes, mime_type: str) -> str:
    """
    Extracts text from a document (PDF or DOCX) provided as bytes.

    Args:
        file_path: The original name of the file (used for logging).
        file_content: The content of the file in bytes.
        mime_type: The MIME type of the file.

    Returns:
        The extracted text as a string, or an empty string if extraction fails or is not supported.
    """
    extracted_text = ""
    try:
        if mime_type == "application/pdf":
            logger.info(f"Attempting to extract text from PDF: {file_path}")
            pdf_reader = PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() or ""
            logger.info(f"Successfully extracted text from PDF: {file_path}. Length: {len(extracted_text)}")
        
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            logger.info(f"Attempting to extract text from DOCX: {file_path}")
            document = docx.Document(io.BytesIO(file_content))
            for para in document.paragraphs:
                extracted_text += para.text + "\n"
            logger.info(f"Successfully extracted text from DOCX: {file_path}. Length: {len(extracted_text)}")
        
        else:
            logger.warning(f"Unsupported MIME type '{mime_type}' for file: {file_path}. No text extraction performed.")
            return ""

    except Exception as e:
        logger.error(f"Error extracting text from file {file_path} (MIME: {mime_type}): {e}", exc_info=True)
        return "" # Return empty string on error

    return extracted_text.strip()

if __name__ == '__main__':
    # Example usage (requires dummy files or manual byte provision)
    # This section is for basic testing and illustration.
    # Proper tests should be in the test suite.

    logger.info("--- Document Parser CLI Test ---")

    # Dummy PDF content (replace with actual bytes from a test PDF if possible)
    # For a real test, you'd create a simple PDF and read its bytes.
    # For now, we'll assume a function to get these bytes.
    try:
        # Create a dummy PDF for testing
        from PyPDF2 import PdfWriter
        pdf_writer = PdfWriter()
        pdf_writer.add_blank_page(width=612, height=792) # Standard US Letter
        # Adding text to a PDF page with PdfWriter directly is not straightforward for content.
        # Usually, you'd use a library like reportlab to generate a PDF with text.
        # For this simple CLI test, we'll rely on unit tests for actual text verification.
        # Here, we'll just test if it attempts to process.
        pdf_bytes_io = io.BytesIO()
        pdf_writer.write(pdf_bytes_io)
        dummy_pdf_bytes = pdf_bytes_io.getvalue()
        
        if dummy_pdf_bytes:
            logger.info("Testing with dummy PDF bytes...")
            text_from_pdf = extract_text_from_document("dummy.pdf", dummy_pdf_bytes, "application/pdf")
            logger.info(f"Extracted from dummy PDF (expected empty for blank): '{text_from_pdf}'")
        else:
            logger.warning("Could not create dummy PDF bytes for CLI test.")

    except ImportError:
        logger.warning("PyPDF2 not fully available for creating a dummy PDF in CLI test.")
    except Exception as e:
        logger.error(f"Error in PDF CLI test section: {e}")


    # Dummy DOCX content
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("Hello, this is a test paragraph in a DOCX file.")
        doc.add_paragraph("Another paragraph for testing.")
        docx_bytes_io = io.BytesIO()
        doc.save(docx_bytes_io)
        dummy_docx_bytes = docx_bytes_io.getvalue()

        if dummy_docx_bytes:
            logger.info("Testing with dummy DOCX bytes...")
            text_from_docx = extract_text_from_document("dummy.docx", dummy_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            logger.info(f"Extracted from dummy DOCX: '{text_from_docx}'")
        else:
            logger.warning("Could not create dummy DOCX bytes for CLI test.")
            
    except ImportError:
        logger.warning("python-docx not available for creating a dummy DOCX in CLI test.")
    except Exception as e:
        logger.error(f"Error in DOCX CLI test section: {e}")

    logger.info("Testing with unsupported file type...")
    text_from_unsupported = extract_text_from_document("dummy.txt", b"some text content", "text/plain")
    logger.info(f"Extracted from unsupported (expected empty): '{text_from_unsupported}'")

    logger.info("--- Document Parser CLI Test Finished ---")
