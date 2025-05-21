import pytest
import io
import logging
from app.services.document_parser import extract_text_from_document

# For creating dummy files for testing
from PyPDF2 import PdfWriter
from docx import Document as DocxDocument

# Configure logging for tests if necessary, or rely on app's logging for coverage
logger = logging.getLogger(__name__)

# --- Fixtures for Dummy File Content ---

@pytest.fixture
def dummy_pdf_content() -> bytes:
    pdf_writer = PdfWriter()
    pdf_writer.add_blank_page(width=612, height=792) # Standard US Letter
    # To add actual text that can be extracted:
    # This is a bit tricky with PyPDF2 directly for simple text.
    # A more robust way involves reportlab or other PDF generation libraries.
    # For a simple test, we'll acknowledge that a blank page has no extractable text.
    # If you have a simple PDF with text, convert it to bytes and use that.
    # For now, let's assume we have a way to add simple text or use a pre-made small PDF.
    # A more advanced approach is to mock the PdfReader.pages part.
    # For this test, we'll use a PDF that *should* have text.
    # Create a PDF with actual text (simplified):
    # This still won't add visible text easily with just PdfWriter for extraction
    # by page.extract_text(). We'll rely on mocking or a pre-existing file for true text.
    # For this example, let's assume the blank page extraction will be an empty string.
    # To properly test actual text extraction, a real PDF with text is needed.
    # Let's adjust this to mock the extraction part if we can't easily create text.

    # Let's try to create a PDF with some text using a method that might work with PdfReader
    # (This is often not reliable for text extraction testing without external tools)
    # For the purpose of this test, we'll assume an empty string from a blank page is expected.
    # And then we'll mock a more complex scenario.
    
    # Simpler approach for testing: create a PDF that *is known* to have certain text.
    # For CI/CD, it's better to generate this or have a static test file.
    # Let's make a PDF with one line of text for this test
    # (Using reportlab would be better for reliable text embedding for tests)

    # Since reportlab is not a listed dependency, we'll make a very simple PDF
    # and accept that `page.extract_text()` might be empty or unreliable
    # for programmatically added text via PyPDF2's lower-level constructs.
    # The test will focus on the flow.

    # For a more reliable test of text extraction, we will mock the PdfReader behavior.
    pdf_bytes_io = io.BytesIO()
    # pdf_writer.write(pdf_bytes_io)
    # For this test, we will provide bytes of a PDF that says "Hello PDF"
    # This typically means having a small, actual PDF file's content.
    # As a placeholder, using a library to create one:
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.drawString(72, 72, "Hello PDF Test Content")
        can.save()
        packet.seek(0)
        return packet.getvalue()
    except ImportError:
        logger.warning("ReportLab not found, PDF text extraction test will be limited.")
        # Fallback to blank PDF if reportlab isn't available
        pdf_writer.add_blank_page(width=612, height=792)
        pdf_writer.write(pdf_bytes_io)
        return pdf_bytes_io.getvalue()


@pytest.fixture
def dummy_docx_content() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Hello DOCX Test Content.")
    doc.add_paragraph("This is the second paragraph.")
    docx_bytes_io = io.BytesIO()
    doc.save(docx_bytes_io)
    docx_bytes_io.seek(0)
    return docx_bytes_io.getvalue()

@pytest.fixture
def corrupted_pdf_content() -> bytes:
    return b"%PDF-1.4\n%corrupted content\n%%EOF"

# --- Test Cases ---

def test_extract_text_from_pdf_success(dummy_pdf_content, caplog):
    caplog.set_level(logging.INFO)
    # If using ReportLab generated PDF:
    expected_text = "Hello PDF Test Content" # Adjust if your generation method differs
    
    # If ReportLab is not available and it's a blank PDF:
    if b"Hello PDF Test Content" not in dummy_pdf_content: # Basic check
         expected_text = "" # Blank page = no text

    extracted_text = extract_text_from_document("test.pdf", dummy_pdf_content, "application/pdf")
    
    # The actual extracted text can be tricky with programmatic PDF generation.
    # We are checking if the function runs and processes.
    # For more precise text matching, a static, well-defined PDF is better.
    if expected_text:
        assert expected_text in extracted_text # Check for presence, formatting might vary
    else:
        assert extracted_text == ""

    assert "Successfully extracted text from PDF: test.pdf" in caplog.text

def test_extract_text_from_docx_success(dummy_docx_content, caplog):
    caplog.set_level(logging.INFO)
    expected_text = "Hello DOCX Test Content.\nThis is the second paragraph."
    extracted_text = extract_text_from_document(
        "test.docx", dummy_docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert expected_text.strip() == extracted_text.strip() # Normalize whitespace
    assert "Successfully extracted text from DOCX: test.docx" in caplog.text

def test_extract_text_unsupported_mime_type(caplog):
    caplog.set_level(logging.WARNING)
    dummy_content = b"Some plain text content"
    extracted_text = extract_text_from_document("test.txt", dummy_content, "text/plain")
    assert extracted_text == ""
    assert "Unsupported MIME type 'text/plain' for file: test.txt" in caplog.text

def test_extract_text_from_corrupted_pdf(corrupted_pdf_content, caplog):
    caplog.set_level(logging.ERROR)
    extracted_text = extract_text_from_document("corrupted.pdf", corrupted_pdf_content, "application/pdf")
    assert extracted_text == ""
    assert "Error extracting text from file corrupted.pdf (MIME: application/pdf)" in caplog.text
    # More specific error from PyPDF2 might be logged, e.g., "PdfReadWarning: EOF marker not found" or similar.

def test_extract_text_from_empty_pdf_content(caplog):
    caplog.set_level(logging.ERROR)
    extracted_text = extract_text_from_document("empty.pdf", b"", "application/pdf")
    assert extracted_text == ""
    assert "Error extracting text from file empty.pdf (MIME: application/pdf)" in caplog.text

def test_extract_text_from_empty_docx_content(caplog):
    # python-docx might handle "empty" bytes differently than a truly corrupted file.
    # It might raise PackageNotFoundError or similar if the bytes don't form a valid ZIP archive (which DOCX is).
    caplog.set_level(logging.ERROR)
    extracted_text = extract_text_from_document(
        "empty.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert extracted_text == ""
    assert "Error extracting text from file empty.docx (MIME: application/vnd.openxmlformats-officedocument.wordprocessingml.document)" in caplog.text

# Example of mocking if generating files with specific text is too complex for some formats
@pytest.mark.skip(reason="Demonstrates mocking, prefer real file content tests if possible")
def test_extract_text_from_pdf_with_mocking(mocker, caplog):
    caplog.set_level(logging.INFO)
    
    # Mock PdfReader and its pages
    mock_pdf_reader_instance = mocker.MagicMock()
    mock_page = mocker.MagicMock()
    mock_page.extract_text.return_value = "Mocked PDF text from page 1. "
    mock_page2 = mocker.MagicMock()
    mock_page2.extract_text.return_value = "Mocked PDF text from page 2."
    mock_pdf_reader_instance.pages = [mock_page, mock_page2]
    
    mocker.patch("app.services.document_parser.PdfReader", return_value=mock_pdf_reader_instance)
    
    dummy_pdf_bytes = b"doesn't matter due to mocking"
    extracted_text = extract_text_from_document("mocked.pdf", dummy_pdf_bytes, "application/pdf")
    
    assert "Mocked PDF text from page 1. Mocked PDF text from page 2." in extracted_text
    assert "Successfully extracted text from PDF: mocked.pdf" in caplog.text
